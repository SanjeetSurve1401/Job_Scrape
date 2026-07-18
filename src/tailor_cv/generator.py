import os
import re
import json
import time
import requests
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from src.tailor_cv.pdf_parser import extract_text_from_pdf
from src.tailor_cv.claude_client import ClaudeClient, accumulate_job_tokens

def sanitize_filename(name: str) -> str:
    clean = re.sub(r'[^a-zA-Z0-9_\-]', '_', name)
    clean = re.sub(r'_+', '_', clean)
    return clean.strip('_').lower()

def get_tailored_cv_data(client, cv_text, job_details) -> dict:
    if hasattr(client, 'get_tailored_cv_data'):
        res = client.get_tailored_cv_data(cv_text, job_details)
        client.last_usage = res.get("usage")
        return res.get("data")
    raise ValueError("Unsupported client type passed to get_tailored_cv_data")

def get_tailored_cl_data(client, cv_text, job_details) -> dict:
    if hasattr(client, 'get_tailored_cl_data'):
        res = client.get_tailored_cl_data(cv_text, job_details)
        client.last_usage = res.get("usage")
        return res.get("data")
    raise ValueError("Unsupported client type passed to get_tailored_cl_data")


def get_tailored_ats_score(client, cv_data: dict, job_details: dict) -> str:
    if hasattr(client, 'get_tailored_ats_score'):
        res = client.get_tailored_ats_score(cv_data, job_details)
        client.last_usage = res.get("usage")
        return res.get("ats_score")
    raise ValueError("Unsupported client type passed to get_tailored_ats_score")

def generate_pdf_cv(cv_data: dict, output_path: str):
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle('CVName', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=22, leading=26, textColor=colors.HexColor('#1A365D'), alignment=1, spaceAfter=4)
    contact_style = ParagraphStyle('CVContact', parent=styles['Normal'], fontSize=9.5, leading=13, textColor=colors.HexColor('#4A5568'), alignment=1, spaceAfter=12)
    section_style = ParagraphStyle('CVSection', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=13, leading=17, textColor=colors.HexColor('#2B6CB0'), spaceBefore=10, spaceAfter=5, keepWithNext=True)
    body_style = ParagraphStyle('CVBody', parent=styles['Normal'], fontSize=9.5, leading=13, textColor=colors.HexColor('#2D3748'), spaceAfter=3)
    bullet_style = ParagraphStyle('CVBullet', parent=styles['Normal'], fontSize=9.5, leading=13, textColor=colors.HexColor('#2D3748'), leftIndent=12, firstLineIndent=-8, spaceAfter=2.5)

    contact = cv_data.get('contact_info', {})
    story.append(Paragraph(contact.get('name', 'Candidate Name'), name_style))
    contact_details = []
    if contact.get('email'): contact_details.append(contact.get('email'))
    if contact.get('phone'): contact_details.append(contact.get('phone'))
    if contact.get('location'): contact_details.append(contact.get('location'))
    if contact.get('linkedin'): contact_details.append(contact.get('linkedin'))
    if contact.get('github'): contact_details.append(contact.get('github'))
    story.append(Paragraph(" | ".join(contact_details), contact_style))
    
    summary = cv_data.get('professional_summary')
    if summary:
        story.append(Paragraph("Professional Summary", section_style))
        story.append(Paragraph(summary, body_style))
        story.append(Spacer(1, 8))
        
    skills = cv_data.get('skills', {})
    if skills:
        story.append(Paragraph("Skills", section_style))
        for skill_type, skill_list in skills.items():
            if skill_list:
                label = skill_type.replace('_', ' ').title()
                story.append(Paragraph(f"<b>{label}:</b> {', '.join(skill_list)}", body_style))
        story.append(Spacer(1, 8))
        
    experience = cv_data.get('work_experience', [])
    if experience:
        story.append(Paragraph("Professional Experience", section_style))
        for job in experience:
            title_line = f"<b>{job.get('role', 'Role')}</b> - {job.get('company', 'Company')}"
            date_line = f"{job.get('start_date', '')} – {job.get('end_date', 'Present')}"
            header_table = Table([[Paragraph(title_line, body_style), Paragraph(date_line, ParagraphStyle('RightText', parent=body_style, alignment=2))]], colWidths=[350, 154])
            header_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('TOPPADDING', (0,0), (-1,-1), 1),
            ]))
            story.append(header_table)
            if job.get('location'):
                story.append(Paragraph(f"<i>{job.get('location')}</i>", ParagraphStyle('ItalicText', parent=body_style, fontSize=8.5, textColor=colors.HexColor('#718096'))))
            for ach in job.get('achievements', []):
                story.append(Paragraph(f"&bull; {ach}", bullet_style))
            story.append(Spacer(1, 4))
            
    projects = cv_data.get('projects', [])
    if projects:
        story.append(Paragraph("Projects", section_style))
        for proj in projects:
            title = proj.get('title', 'Project')
            tech = proj.get('technologies', '')
            proj_header = f"<b>{title}</b>"
            if tech:
                proj_header += f" (<i>Technologies: {tech}</i>)"
            story.append(Paragraph(proj_header, body_style))
            story.append(Paragraph(proj.get('description', ''), body_style))
            story.append(Spacer(1, 3))
            
    education = cv_data.get('education', [])
    if education:
        story.append(Paragraph("Education", section_style))
        for edu in education:
            degree_str = f"<b>{edu.get('degree', 'Degree')}</b> in {edu.get('major', 'Major')}"
            edu_table = Table([[Paragraph(degree_str, body_style), Paragraph(edu.get('graduation_year', ''), ParagraphStyle('RightText', parent=body_style, alignment=2))]], colWidths=[400, 104])
            edu_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('TOPPADDING', (0,0), (-1,-1), 1),
            ]))
            story.append(edu_table)
            story.append(Paragraph(edu.get('institution', ''), ParagraphStyle('SubText', parent=body_style, textColor=colors.HexColor('#718096'))))
            story.append(Spacer(1, 3))
            
    certs = cv_data.get('certifications', [])
    if certs:
        story.append(Paragraph("Certifications", section_style))
        for cert in certs:
            story.append(Paragraph(f"&bull; {cert}", bullet_style))
            
    doc.build(story)

def generate_pdf_cover_letter(cl_data: dict, output_path: str):
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    styles = getSampleStyleSheet()
    
    body_style = ParagraphStyle('CLBody', parent=styles['Normal'], fontSize=11, leading=15, textColor=colors.HexColor('#2D3748'), spaceAfter=10)
    sender_style = ParagraphStyle('CLSender', parent=styles['Normal'], fontSize=11, leading=15, textColor=colors.HexColor('#1A365D'), spaceAfter=3)
    date_style = ParagraphStyle('CLDate', parent=styles['Normal'], fontSize=11, leading=15, textColor=colors.HexColor('#4A5568'), spaceAfter=12)
    recipient_style = ParagraphStyle('CLRecipient', parent=styles['Normal'], fontSize=11, leading=15, textColor=colors.HexColor('#2D3748'), spaceAfter=10)
    subject_style = ParagraphStyle('CLSubject', parent=styles['Normal'], fontSize=11, leading=15, fontName='Helvetica-Bold', textColor=colors.HexColor('#1A365D'), spaceAfter=12)

    story.append(Paragraph(cl_data.get('date', 'July 6, 2026'), date_style))
    story.append(Paragraph(cl_data.get('recipient_name', 'Hiring Manager / Recruiter'), recipient_style))
    story.append(Paragraph(cl_data.get('company_name', 'Company Name'), recipient_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph(cl_data.get('subject', 'Application'), subject_style))
    story.append(Paragraph(cl_data.get('salutation', 'Dear Hiring Manager,'), body_style))
    story.append(Paragraph(cl_data.get('opening_paragraph', ''), body_style))
    for para in cl_data.get('body_paragraphs', []):
        story.append(Paragraph(para, body_style))
    story.append(Paragraph(cl_data.get('closing_paragraph', ''), body_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph(cl_data.get('sign_off', 'Sincerely,'), body_style))
    story.append(Spacer(1, 4))
    story.append(Paragraph(f"<b>{cl_data.get('sender_name', 'Candidate Name')}</b>", sender_style))
    
    doc.build(story)

def generate_docx_cv(cv_data: dict, output_path: str):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    navy_blue = RGBColor(0x1A, 0x36, 0x5D)
    light_blue = RGBColor(0x2B, 0x6C, 0xB0)
    charcoal = RGBColor(0x2D, 0x37, 0x48)
    grey = RGBColor(0x71, 0x80, 0x96)
    
    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = light_blue
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.keep_with_next = True

    contact = cv_data.get('contact_info', {})
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(contact.get('name', 'Candidate Name'))
    run_name.bold = True
    run_name.font.size = Pt(22)
    run_name.font.color.rgb = navy_blue
    p_name.paragraph_format.space_after = Pt(4)
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_details = []
    if contact.get('email'): contact_details.append(contact.get('email'))
    if contact.get('phone'): contact_details.append(contact.get('phone'))
    if contact.get('location'): contact_details.append(contact.get('location'))
    if contact.get('linkedin'): contact_details.append(contact.get('linkedin'))
    if contact.get('github'): contact_details.append(contact.get('github'))
    run_contact = p_contact.add_run(" | ".join(contact_details))
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = charcoal
    p_contact.paragraph_format.space_after = Pt(10)
    
    summary = cv_data.get('professional_summary')
    if summary:
        add_section_heading("Professional Summary")
        p = doc.add_paragraph()
        run = p.add_run(summary)
        run.font.size = Pt(9.5)
        run.font.color.rgb = charcoal
        p.paragraph_format.space_after = Pt(6)
        
    skills = cv_data.get('skills', {})
    if skills:
        add_section_heading("Skills")
        for skill_type, skill_list in skills.items():
            if skill_list:
                label = skill_type.replace('_', ' ').title()
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                r_lbl = p.add_run(f"{label}: ")
                r_lbl.bold = True
                r_lbl.font.size = Pt(9.5)
                r_val = p.add_run(", ".join(skill_list))
                r_val.font.size = Pt(9.5)
                r_val.font.color.rgb = charcoal
                
    experience = cv_data.get('work_experience', [])
    if experience:
        add_section_heading("Professional Experience")
        for job in experience:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(5.0)
            table.columns[1].width = Inches(2.0)
            cell_left = table.cell(0, 0)
            cell_right = table.cell(0, 1)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(0)
            r_title = p_left.add_run(f"{job.get('role', 'Role')} at {job.get('company', 'Company')}")
            r_title.bold = True
            r_title.font.size = Pt(9.5)
            r_title.font.color.rgb = charcoal
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(0)
            r_date = p_right.add_run(f"{job.get('start_date', '')} – {job.get('end_date', 'Present')}")
            r_date.font.size = Pt(9.5)
            r_date.font.color.rgb = charcoal
            
            if job.get('location'):
                p_loc = doc.add_paragraph()
                p_loc.paragraph_format.space_after = Pt(2)
                r_loc = p_loc.add_run(job.get('location'))
                r_loc.italic = True
                r_loc.font.size = Pt(8.5)
                r_loc.font.color.rgb = grey
                
            for ach in job.get('achievements', []):
                p_ach = doc.add_paragraph(style='List Bullet')
                p_ach.paragraph_format.space_after = Pt(2.5)
                p_ach.paragraph_format.left_indent = Inches(0.2)
                r_ach = p_ach.add_run(ach)
                r_ach.font.size = Pt(9.5)
                r_ach.font.color.rgb = charcoal
                
    projects = cv_data.get('projects', [])
    if projects:
        add_section_heading("Projects")
        for proj in projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r_title = p.add_run(proj.get('title', 'Project'))
            r_title.bold = True
            r_title.font.size = Pt(9.5)
            r_title.font.color.rgb = charcoal
            if proj.get('technologies'):
                r_tech = p.add_run(f" (Technologies: {proj.get('technologies')})")
                r_tech.italic = True
                r_tech.font.size = Pt(8.5)
                r_tech.font.color.rgb = grey
            p_desc = doc.add_paragraph()
            p_desc.paragraph_format.space_after = Pt(3)
            r_desc = p_desc.add_run(proj.get('description', ''))
            r_desc.font.size = Pt(9.5)
            r_desc.font.color.rgb = charcoal
            
    education = cv_data.get('education', [])
    if education:
        add_section_heading("Education")
        for edu in education:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(5.5)
            table.columns[1].width = Inches(1.5)
            cell_left = table.cell(0, 0)
            cell_right = table.cell(0, 1)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(0)
            r_degree = p_left.add_run(f"{edu.get('degree', 'Degree')} in {edu.get('major', 'Major')}")
            r_degree.bold = True
            r_degree.font.size = Pt(9.5)
            r_degree.font.color.rgb = charcoal
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(0)
            r_grad = p_right.add_run(edu.get('graduation_year', ''))
            r_grad.font.size = Pt(9.5)
            r_grad.font.color.rgb = charcoal
            p_inst = doc.add_paragraph()
            p_inst.paragraph_format.space_after = Pt(3)
            r_inst = p_inst.add_run(edu.get('institution', ''))
            r_inst.font.size = Pt(9.0)
            r_inst.font.color.rgb = grey
            
    certs = cv_data.get('certifications', [])
    if certs:
        add_section_heading("Certifications")
        for cert in certs:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2.5)
            p.paragraph_format.left_indent = Inches(0.2)
            r_cert = p.add_run(cert)
            r_cert.font.size = Pt(9.5)
            r_cert.font.color.rgb = charcoal
    doc.save(output_path)

def generate_docx_cover_letter(cl_data: dict, output_path: str):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    charcoal = RGBColor(0x2D, 0x37, 0x48)
    navy_blue = RGBColor(0x1A, 0x36, 0x5D)
    
    def add_p(text, font_size=11, bold=False, space_after=10, italic=False, color=charcoal):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        run.font.color.rgb = color
        return p

    add_p(cl_data.get('date', 'July 6, 2026'), space_after=15)
    add_p(cl_data.get('recipient_name', 'Hiring Manager / Recruiter'), space_after=2)
    add_p(cl_data.get('company_name', 'Company Name'), space_after=10)
    add_p(cl_data.get('subject', 'Application'), bold=True, color=navy_blue, space_after=12)
    add_p(cl_data.get('salutation', 'Dear Hiring Manager,'), space_after=10)
    add_p(cl_data.get('opening_paragraph', ''), space_after=10)
    for para in cl_data.get('body_paragraphs', []):
        add_p(para, space_after=10)
    add_p(cl_data.get('closing_paragraph', ''), space_after=10)
    add_p(cl_data.get('sign_off', 'Sincerely,'), space_after=5)
    add_p(cl_data.get('sender_name', 'Candidate Name'), bold=True, color=navy_blue, space_after=0)
    doc.save(output_path)

def generate_tailored_documents(cv_path: str, job: dict, model: str, output_base_dir: str = "outputs"):
    try:
        from src.tailor_cv.pdf_parser import extract_text_from_pdf
        cv_text = extract_text_from_pdf(cv_path)
        
        from src.tailor_cv.claude_client import ClaudeClient, accumulate_job_tokens
        client = ClaudeClient(model=model)
        
        print(f"  \n[Tailoring CV] Calling LLM (Claude) for '{job.get('title')}' at '{job.get('company')}'...\n")
        cv_data = get_tailored_cv_data(client, cv_text, job)
        if hasattr(client, 'last_usage') and client.last_usage:
            accumulate_job_tokens(job, client.last_usage)
        
        time.sleep(2.0)
        
        print(f"  [Tailoring Cover Letter] Calling LLM (Claude) for '{job.get('title')}' at '{job.get('company')}'...")
        cl_data = get_tailored_cl_data(client, cv_text, job)
        if hasattr(client, 'last_usage') and client.last_usage:
            accumulate_job_tokens(job, client.last_usage)
        
        company_clean = sanitize_filename(job.get('company', 'Unknown_Company'))
        title_clean = sanitize_filename(job.get('title', 'Unknown_Title'))
        folder_name = f"{company_clean}_{title_clean}"
        target_dir = os.path.join(output_base_dir, folder_name)
        os.makedirs(target_dir, exist_ok=True)
        
        cv_pdf_path = os.path.join(target_dir, "tailored_cv.pdf")
        cv_docx_path = os.path.join(target_dir, "tailored_cv.docx")
        cl_pdf_path = os.path.join(target_dir, "cover_letter.pdf")
        cl_docx_path = os.path.join(target_dir, "cover_letter.docx")
        
        generate_pdf_cv(cv_data, cv_pdf_path)
        generate_docx_cv(cv_data, cv_docx_path)
        generate_pdf_cover_letter(cl_data, cl_pdf_path)
        generate_docx_cover_letter(cl_data, cl_docx_path)
        
        # Get tailored ATS score
        print(f"  \n[ATS Scanner] Simulating ATS match score for tailored CV...\n")
        ats_score = get_tailored_ats_score(client, cv_data, job)
        job["ats_score"] = ats_score
        if hasattr(client, 'last_usage') and client.last_usage:
            accumulate_job_tokens(job, client.last_usage)
        
        # Save job details as JSON in the same folder
        job_json_path = os.path.join(target_dir, "job_details.json")
        with open(job_json_path, 'w', encoding='utf-8') as jf:
            json.dump(job, jf, indent=2, ensure_ascii=False)
            
        print(f"  \n[Completed] Saved tailored CV, Cover Letter, and Job Details in: {target_dir} (ATS Score: {ats_score})\n")
        return ats_score
    except Exception as e:
        print(f"  [Error] Failed to generate tailored documents for '{job.get('title')}' at '{job.get('company')}': {e}")
        return None
